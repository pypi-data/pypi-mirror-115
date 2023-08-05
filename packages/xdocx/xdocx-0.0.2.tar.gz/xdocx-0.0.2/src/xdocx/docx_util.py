from openpyxl import load_workbook
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


class get_xlsx_data:
    VAILD_ROW = 4
    def __init__(self,file_path):
        try:
            print(file_path)
            wb = load_workbook(filename = file_path)
            self.__testcase_handle = wb['测试用例列表']
        except:
            print("parse xlsx file error.")
            exit(-1)

    def get_total_number(self,table):
        ret = 0
        for _i in table:
            ret += 1
        return ret

    def get_all_testcase(self):
        ret_data = []
        table_num = self.get_total_number(self.__testcase_handle)
        for _i in range(table_num - get_xlsx_data.VAILD_ROW + 1):
            ret_data.append([self.__testcase_handle[4 + _i][0].value,self.__testcase_handle[4 + _i][2].value,self.__testcase_handle[4 + _i][3].value])
        return ret_data



class write_test_record:
    VAILD_DATA_ROWS = 2
    def __init__(self,input_file,output_file):
        self.__output_file = output_file
        try:
            self.__docx_handle = Document(input_file)
            self.__record_table = self.__docx_handle.tables[0]


        except:
            print("parse input file error.")
            exit(-1)
    
    def add_table_rows(self,num):
        for i in range(num):
            cell = self.__record_table.add_row()
    

    def write_record(self,data):
        total_testcase_num = len(data)
        self.add_table_rows(total_testcase_num - 1)
        index = 0 
        standard_format = self.__record_table.rows[write_test_record.VAILD_DATA_ROWS]

        for _i in self.__record_table.rows[write_test_record.VAILD_DATA_ROWS:len(data) + write_test_record.VAILD_DATA_ROWS]:
            _i.cells[0].text = data[index][0]
            _i.cells[1].text = data[index][1]
            _i.cells[2].text = data[index][2]
            index += 1
        self.__docx_handle.save(self.__output_file)